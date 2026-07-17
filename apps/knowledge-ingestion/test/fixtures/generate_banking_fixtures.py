#!/usr/bin/env python3
"""
Generate synthetic co-operative-bank (UCB) knowledge fixtures for W2 parser tests.

ALL CONTENT IS FABRICATED. No real bank, customer, RBI circular, or CMDB data is
used (per docs/ai-copilot/TESTING_STRATEGY.md — synthetic only). The fixtures are
realistic in *shape* (correct CI types, service names, document structure) so the
parsers and chunker are exercised against co-op-bank-realistic inputs (ADR-003),
while the pipeline code itself stays fully vertical-agnostic.

Outputs into ./banking/:
  rbi-style-circular.pdf      ~6-page synthetic circular (numbered headings)
  rbi-circular-20page.pdf     20-page padded circular (CP2.2 chunker fixture)
  cbs-eod-runbook.docx        CBS end-of-day runbook with Heading 1/2/3 hierarchy
  cmdb-export.xlsx            UCB CMDB export (incl. sponsor_bank_link / npci_link)
  branch-topology.pdf         hub-and-spoke branch diagram (Visio-export shape)
  upi-recon-sop.md            UPI reconciliation SOP (markdown headings)
  atm-cashout-sop.txt         ATM cash-out SOP (plain text)

Run:  python3 generate_banking_fixtures.py
"""
import os

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfgen import canvas
import docx
from docx.shared import Pt
import openpyxl

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "banking")
os.makedirs(OUT, exist_ok=True)

SYNTHETIC_NOTE = (
    "SYNTHETIC FIXTURE — fabricated for Canaris AI Copilot W2 testing. "
    "Not a real RBI circular and not real bank data."
)


# --------------------------------------------------------------------------- #
# 1 + 6. RBI-style circular PDF (numbered headings + sections)
# --------------------------------------------------------------------------- #
CIRCULAR_SECTIONS = [
    ("1 Introduction",
     "This circular sets out the baseline operational resilience expectations for "
     "the core banking and digital payment systems operated by the bank. It is "
     "issued for internal operational reference only and is entirely synthetic."),
    ("1.1 Applicability",
     "These expectations apply to all customer-facing transactional services, "
     "including core banking, UPI/IMPS retail payments, NEFT/RTGS interbank "
     "transfers, and ATM/card switching, whether operated on-premises or as a "
     "hosted/shared service."),
    ("1.2 Scope",
     "The scope covers availability monitoring, incident reporting timelines, "
     "business-continuity posture, and the change-control discipline required for "
     "tier-1 services. Connectivity to the sponsor commercial bank and to the "
     "national payment infrastructure is explicitly in scope."),
    ("2 Operational Resilience",
     "The bank shall maintain continuous monitoring of all tier-1 services and "
     "their supporting configuration items, including the sponsor-bank link and "
     "the NPCI link through which retail payment settlement flows."),
    ("2.1 Availability Targets",
     "Retail digital payment rails operate on a 24x7 basis and shall be monitored "
     "accordingly. Branch-hours services may follow branch operating windows."),
    ("2.2 Dependency Mapping",
     "Each tier-1 service shall have a documented dependency map identifying the "
     "configuration items on which it depends, so that the operational impact of "
     "any single component failure can be assessed promptly."),
    ("3 Incident Management",
     "Material operational incidents affecting customer transactions shall be "
     "triaged, escalated, and recorded. The reporting window and the precise "
     "escalation matrix are defined in the bank's internal incident SOP."),
    ("3.1 Classification",
     "Incidents shall be classified by customer impact and by the tier of the "
     "affected service. A failure of the sponsor-bank link that blocks UPI/IMPS "
     "settlement is, by definition, a tier-1 incident."),
    ("3.2 Reporting Timeline",
     "Reportable incidents shall be notified to the relevant authority within the "
     "stipulated window. (Exact hour-counts are deliberately omitted from this "
     "synthetic fixture and must be confirmed against current regulation.)"),
    ("4 Business Continuity",
     "Each tier-1 service shall have a defined disaster-recovery node and a tested "
     "failover procedure. Recovery objectives shall be documented and exercised "
     "periodically as part of the BCP programme."),
    ("4.1 Disaster Recovery",
     "A tier-1 service with no mapped DR node is considered to violate the expected "
     "business-continuity posture and shall be flagged for remediation."),
    ("5 Conclusion",
     "Adherence to these baseline expectations is intended to keep customer-facing "
     "services resilient. This document is synthetic and for test use only."),
]


def _circular_story(styles, repeats=1):
    title_style = ParagraphStyle(
        "CircTitle", parent=styles["Title"], fontSize=16, spaceAfter=6
    )
    h_style = ParagraphStyle(
        "CircHeading", parent=styles["Heading2"], fontSize=12, spaceBefore=10,
        spaceAfter=4,
    )
    body = ParagraphStyle("CircBody", parent=styles["BodyText"], fontSize=10,
                          leading=14)
    note = ParagraphStyle("CircNote", parent=styles["Italic"], fontSize=8,
                          textColor="#888888")
    story = []
    story.append(Paragraph("Master Circular on Operational Resilience of "
                           "Core Banking and Payment Systems", title_style))
    story.append(Paragraph("Ref: SYN/UCB/OPS/2026-01 (synthetic)", note))
    story.append(Paragraph(SYNTHETIC_NOTE, note))
    story.append(Spacer(1, 6 * mm))
    for r in range(repeats):
        for heading, text in CIRCULAR_SECTIONS:
            h = heading if r == 0 else f"{heading} (cont. {r + 1})"
            story.append(Paragraph(h, h_style))
            story.append(Paragraph(text, body))
        if r != repeats - 1:
            story.append(PageBreak())
    return story


def make_circular(path, repeats, target_pages=None):
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=20 * mm, rightMargin=20 * mm,
                            topMargin=18 * mm, bottomMargin=18 * mm,
                            title="Synthetic UCB Operational Resilience Circular")
    doc.build(_circular_story(styles, repeats=repeats))


# --------------------------------------------------------------------------- #
# 2. CBS end-of-day runbook DOCX (heading hierarchy)
# --------------------------------------------------------------------------- #
def make_cbs_runbook(path):
    d = docx.Document()
    d.add_heading("CBS End-of-Day (EOD) Runbook", level=0)
    p = d.add_paragraph(SYNTHETIC_NOTE)
    p.runs[0].italic = True

    d.add_heading("1. Purpose", level=1)
    d.add_paragraph(
        "This runbook describes the synthetic end-of-day batch procedure for the "
        "core banking system (CBS) of a medium urban co-operative bank. It is a "
        "fabricated training artifact and does not describe any real system.")

    d.add_heading("2. Pre-EOD Checklist", level=1)
    d.add_heading("2.1 Connectivity Verification", level=2)
    d.add_paragraph(
        "Confirm that the sponsor-bank link and the NPCI link are healthy and that "
        "all branch WAN tunnels are up before initiating the batch.")
    d.add_heading("2.2 Pending Transaction Drain", level=2)
    d.add_paragraph(
        "Ensure all in-flight UPI/IMPS and NEFT/RTGS transactions have settled or "
        "been queued for the next cycle. Verify the cheque-clearing (CTS) queue is "
        "drained.")

    d.add_heading("3. EOD Execution", level=1)
    d.add_heading("3.1 Interest and GL Posting", level=2)
    d.add_paragraph(
        "Trigger interest accrual and general-ledger posting. Monitor the CBS "
        "database server for lock contention during the posting window.")
    d.add_heading("3.2 Day Rollover", level=2)
    d.add_paragraph(
        "Advance the CBS business date. Confirm the new working date on the CBS "
        "application server before re-enabling channels.")

    d.add_heading("4. Post-EOD Validation", level=1)
    d.add_paragraph(
        "Reconcile control totals, confirm DR replication caught up to the DR site "
        "node, and re-open customer channels. Record completion in the operations "
        "log.")

    d.add_heading("5. Rollback", level=1)
    d.add_paragraph(
        "If posting fails midway, halt the batch, restore from the pre-EOD "
        "snapshot, and escalate per the incident SOP. Do not re-open channels until "
        "control totals reconcile.")
    d.save(path)


# --------------------------------------------------------------------------- #
# 3. CMDB export XLSX (table-aware; co-op-specific CI types)
# --------------------------------------------------------------------------- #
CMDB_COLUMNS = [
    "ci_id", "ci_name", "ci_type", "criticality_tier",
    "business_service", "technical_owner", "location",
]
CMDB_ROWS = [
    ["CI-0001", "CBS App Node 1", "cbs_application_server", "tier_1",
     "core_banking", "core.ops@synthbank.example", "DC-Primary"],
    ["CI-0002", "CBS DB Node 1", "cbs_database_server", "tier_1",
     "core_banking", "dba.team@synthbank.example", "DC-Primary"],
    ["CI-0003", "CBS Hosted Service", "cbs_hosted_service", "tier_1",
     "core_banking", "vendor.mgmt@synthbank.example", "Hosted-Shared"],
    ["CI-0004", "UPI Switch 1", "upi_switch", "tier_1",
     "upi_imps", "payments.ops@synthbank.example", "DC-Primary"],
    ["CI-0005", "Sponsor Bank Link A", "sponsor_bank_link", "tier_1",
     "upi_imps", "payments.ops@synthbank.example", "WAN-Edge"],
    ["CI-0006", "NPCI Link A", "npci_link", "tier_1",
     "upi_imps", "payments.ops@synthbank.example", "WAN-Edge"],
    ["CI-0007", "Payment Gateway 1", "payment_gateway", "tier_1",
     "upi_imps", "payments.ops@synthbank.example", "DC-Primary"],
    ["CI-0008", "ATM Switch 1", "atm_switch", "tier_1",
     "atm_card_services", "channels.ops@synthbank.example", "DC-Primary"],
    ["CI-0009", "HSM Device 1", "hsm_device", "tier_1",
     "atm_card_services", "security.ops@synthbank.example", "DC-Primary"],
    ["CI-0010", "Sponsor Bank Link B (DR)", "sponsor_bank_link", "tier_1",
     "neft_rtgs", "payments.ops@synthbank.example", "DC-DR"],
    ["CI-0011", "Core Router 1", "core_router", "tier_2",
     "branch_wan", "network.ops@synthbank.example", "DC-Primary"],
    ["CI-0012", "Branch Router BR-014", "branch_router", "tier_2",
     "branch_wan", "network.ops@synthbank.example", "Branch-014"],
    ["CI-0013", "Firewall Edge 1", "firewall", "tier_2",
     "branch_wan", "security.ops@synthbank.example", "WAN-Edge"],
    ["CI-0014", "DR Site Node 1", "dr_site_node", "tier_1",
     "core_banking", "bcp.team@synthbank.example", "DC-DR"],
    ["CI-0015", "Backup System 1", "backup_system", "tier_3",
     "document_management", "infra.ops@synthbank.example", "DC-Primary"],
    ["CI-0016", "Regulatory Reporting Server", "cbs_application_server", "tier_2",
     "regulatory_reporting", "compliance.ops@synthbank.example", "DC-Primary"],
]


def make_cmdb_xlsx(path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "cmdb"
    ws.append(CMDB_COLUMNS)
    for row in CMDB_ROWS:
        ws.append(row)
    wb.save(path)


def make_cmdb_csv(path):
    import csv as _csv
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = _csv.writer(f)
        w.writerow(CMDB_COLUMNS)
        w.writerows(CMDB_ROWS)


# --------------------------------------------------------------------------- #
# 4. Branch topology PDF (hub-and-spoke; Visio-export shape, sparse text)
# --------------------------------------------------------------------------- #
def make_branch_topology(path):
    c = canvas.Canvas(path, pagesize=A4)
    w, h = A4
    cx, cy = w / 2, h / 2
    c.setFont("Helvetica", 8)
    c.drawString(15 * mm, h - 12 * mm, SYNTHETIC_NOTE)
    c.setFont("Helvetica-Bold", 10)
    c.drawCentredString(cx, h - 22 * mm, "Branch Network Topology (synthetic, hub-and-spoke)")

    # Hub box (data centre)
    c.setLineWidth(1.5)
    c.rect(cx - 28 * mm, cy - 10 * mm, 56 * mm, 20 * mm)
    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(cx, cy + 1 * mm, "DC-Primary Hub")
    c.setFont("Helvetica", 7)
    c.drawCentredString(cx, cy - 5 * mm, "core_router / firewall")

    # Spoke branch boxes around the hub
    import math
    spokes = [
        "Branch-014", "Branch-021", "Branch-007", "Branch-033",
        "Branch-002", "Branch-045", "DC-DR", "Sponsor-Link",
    ]
    radius = 70 * mm
    for i, label in enumerate(spokes):
        ang = (2 * math.pi * i) / len(spokes)
        sx = cx + radius * math.cos(ang)
        sy = cy + radius * math.sin(ang)
        c.setLineWidth(0.7)
        c.line(cx, cy, sx, sy)
        c.rect(sx - 16 * mm, sy - 6 * mm, 32 * mm, 12 * mm)
        c.setFont("Helvetica", 7)
        c.drawCentredString(sx, sy - 1 * mm, label)
    c.showPage()
    c.save()


# --------------------------------------------------------------------------- #
# 5. UPI reconciliation SOP (markdown) + ATM cash-out SOP (txt)
# --------------------------------------------------------------------------- #
UPI_RECON_MD = """\
# UPI Reconciliation SOP

> SYNTHETIC FIXTURE — fabricated for Canaris AI Copilot W2 testing.

## 1. Purpose

This SOP describes the daily reconciliation of UPI/IMPS retail payment
transactions between the bank's switch, the sponsor bank, and NPCI. All
amounts, IDs, and timings here are fabricated.

## 2. Inputs

- Switch transaction log (UPI switch)
- Sponsor-bank settlement file
- NPCI raw data file

### 2.1 File Availability

Confirm all three inputs are available before starting. A missing sponsor-bank
settlement file blocks reconciliation and must be escalated.

## 3. Procedure

### 3.1 Match

Match each switch transaction against the NPCI raw data and the sponsor-bank
settlement on the transaction reference.

### 3.2 Investigate Exceptions

Unmatched or amount-mismatched entries are exceptions. Tag each exception with a
reason code and route to the payments operations queue.

## 4. Sign-off

Reconciliation is signed off only when all exceptions are resolved or formally
carried forward. Record the outcome in the operations log.
"""

ATM_CASHOUT_TXT = """\
ATM CASH-OUT SOP (SYNTHETIC FIXTURE — fabricated for W2 testing)

1. PURPOSE
   Procedure for responding to an ATM cash-out (cash-depletion) condition across
   the synthetic branch ATM fleet. All details are fabricated.

2. DETECTION
   The ATM switch raises a low-cash or cash-out signal for a terminal. The
   channels operations team receives the alert and verifies it.

3. RESPONSE
   3.1 Confirm the terminal cash balance from the switch.
   3.2 Raise a cash-replenishment request to the cash-management vendor.
   3.3 If the terminal serves a high-footfall branch, prioritise replenishment.

4. ESCALATION
   If multiple terminals in one region go cash-out simultaneously, escalate to
   the duty manager as a potential cash-logistics incident.

5. CLOSURE
   Confirm replenishment, return the terminal to service, and record the event
   in the operations log.
"""


def write_text(path, content):
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)


def main():
    make_circular(os.path.join(OUT, "rbi-style-circular.pdf"), repeats=3)
    make_circular(os.path.join(OUT, "rbi-circular-20page.pdf"), repeats=10)
    make_cbs_runbook(os.path.join(OUT, "cbs-eod-runbook.docx"))
    make_cmdb_xlsx(os.path.join(OUT, "cmdb-export.xlsx"))
    make_cmdb_csv(os.path.join(OUT, "cmdb-export.csv"))
    make_branch_topology(os.path.join(OUT, "branch-topology.pdf"))
    write_text(os.path.join(OUT, "upi-recon-sop.md"), UPI_RECON_MD)
    write_text(os.path.join(OUT, "atm-cashout-sop.txt"), ATM_CASHOUT_TXT)
    for fn in sorted(os.listdir(OUT)):
        full = os.path.join(OUT, fn)
        print(f"  {fn:32s} {os.path.getsize(full):>8d} bytes")


if __name__ == "__main__":
    main()
